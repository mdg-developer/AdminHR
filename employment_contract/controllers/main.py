from odoo import http
from odoo.http import request
from docx import Document
import zipfile
import io
from datetime import datetime, date, time, timedelta
from odoo.modules.module import get_module_resource
import base64


class Binary(http.Controller):

    @http.route('/web/binary/download_docx_report/<int:employee>/<string:employee_name>', type='http', auth="public")
    def download_function_descriptions(self, employee=None, employee_name=None):

        if employee:
            res = request.env['hr.employee'].get_employee_info(employee)
            stream = io.BytesIO()
            # import pdb
            # pdb.set_trace() 
            attachment = request.env['ir.attachment'].sudo().search(
                [('res_model', '=', 'contract.config'), ('name', '=', 'Employment Contract')], order="id desc", limit=1)

            contract_file_path = attachment._full_path(attachment.store_fname)
            # request.env['contract.config'].search([('employee','=',employee)]).contract_file
            # contract_file_path = get_module_resource('employment_contract', 'data', 'EC_WB_HO.docx')
            f = open(contract_file_path, 'rb')
            doc = Document(f)
            for p in doc.paragraphs:
                if len(p.text) > 0 and p.text.find("#") != -1:
                    inline = p.runs
                    for i in range(len(inline)):
                        line = p.text
                        text = inline[i].text
                        words = text.split("#")
                        new_words = ''
                        new_line = ''
                        for j in range(len(words)):
                            word = words[j]
                            if word in res:
                                word = res[word]
                            new_words += word
                        inline[i].text = new_words
            doc.save('/tmp/EC_WB_HO.docx')
            f.close()

            contract_doc = open('/tmp/EC_WB_HO.docx', 'rb').read();
            open('/tmp/EC_WB_HO.docx', 'wb').write(contract_doc)
            return request.make_response(contract_doc, [('Content-Type', 'application/msword'),
                                                        ('Content-Disposition', 'attachment')])

    @http.route('/web/binary/download_docx_template_report/<int:employee>/<string:employee_name>/<string:doc_name>',
                type='http', auth="public")
    def download_doc_template_function_descriptions(self, employee=None, employee_name=None, doc_name=None):

        if employee:
            res = request.env['hr.employee'].get_employee_info(employee)
            stream = io.BytesIO()
            # import pdb
            # pdb.set_trace()
            attachment = request.env['ir.attachment'].sudo().search(
                [('res_model', '=', 'document.template.config'), ('name', '=', doc_name)], order="id desc", limit=1)
            doc_template_id = request.env['document.template.config'].sudo().search([('name', '=', doc_name)],
                                                                                    order="id desc", limit=1)
            contract_file_path = attachment._full_path(attachment.store_fname)
            # request.env['contract.config'].search([('employee','=',employee)]).contract_file
            # contract_file_path = get_module_resource('employment_contract', 'data', 'EC_WB_HO.docx')
            f = open(contract_file_path, 'rb')
            doc = Document(f)
            for p in doc.paragraphs:
                if len(p.text) > 0 and p.text.find("#") != -1:
                    inline = p.runs
                    for i in range(len(inline)):
                        line = p.text
                        text = inline[i].text
                        words = text.split("#")
                        new_words = ''
                        new_line = ''
                        get_word = False
                        for j in range(len(words)):
                            word = words[j]

                            if doc_template_id and not get_word:
                                emp = request.env['hr.employee'].sudo().search([('id', '=', employee)], order="id desc",
                                                                               limit=1)
                                for line_word in doc_template_id.template_lines:
                                    if word == 'today':
                                        today_date = datetime.now() + timedelta(hours=+6, minutes=+30)
                                        word = today_date.strftime('%d/%m/%Y')
                                        get_word = True
                                        break
                                    elif line_word.model_field.ttype == 'many2one' and line_word.model_field.name == word:
                                        dummy_id = emp
                                        tmp_id = dummy_id.read([line_word.model_field.name])[0]
                                        if tmp_id:
                                            val_id = tmp_id[line_word.model_field.name][
                                                1]  # self.env[line_word.model_name.model].sudo().search([('id', '=', tmp_id)])
                                            word = val_id
                                            get_word = True
                                        break
                                    elif line_word.model_field.name == word:
                                        dummy_id = emp
                                        tmp_id = dummy_id.read([line_word.model_field.name])[0]
                                        if tmp_id:
                                            word = tmp_id[line_word.model_field.name]
                                            get_word = True
                                        break

                            new_words += word
                        inline[i].text = new_words
            doc_save_path = '/tmp/' + doc_name + ".docx"
            doc.save(doc_save_path)
            f.close()

            contract_doc = open(doc_save_path, 'rb').read();
            open(doc_save_path, 'wb').write(contract_doc)
            return request.make_response(contract_doc, [('Content-Type', 'application/msword'),
                                                        ('Content-Disposition', 'attachment')])
